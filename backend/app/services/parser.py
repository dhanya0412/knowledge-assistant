import os
from pathlib import Path


SUPPORTED_EXTENSIONS = {"txt", "pdf", "docx"}


class ParserError(Exception):
    """Base exception for parser failures."""


class UnsupportedFileTypeError(ParserError):
    """Raised when a file extension is not supported."""


class MissingFileError(ParserError):
    """Raised when the target file is missing or is not a file."""


class EmptyDocumentError(ParserError):
    """Raised when no extractable text is found."""


def extract_text(filepath, filename=None):
    path = Path(filepath)
    if not path.exists() or not path.is_file():
        raise MissingFileError("file does not exist")

    extension = _detect_extension(path, filename)

    if extension == "txt":
        text = _extract_txt_text(path)
    elif extension == "pdf":
        text = _extract_pdf_text(path)
    elif extension == "docx":
        text = _extract_docx_text(path)
    else:
        raise UnsupportedFileTypeError("file type is not supported")

    if not text or not text.strip():
        raise EmptyDocumentError("no extractable text found")

    return text


def _detect_extension(path, filename=None):
    source_name = filename or path.name
    extension = os.path.splitext(source_name)[1].lower().lstrip(".")

    if not extension:
        extension = path.suffix.lower().lstrip(".")

    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError("file type is not supported")

    return extension


def _extract_txt_text(path):
    content = path.read_bytes()

    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    return content.decode("utf-8", errors="replace")


def _extract_pdf_text(path):
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ParserError("pypdf dependency is not installed") from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise ParserError("failed to read pdf file") from exc

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ParserError("pdf file is encrypted") from exc

    page_text = []
    for page in reader.pages:
        try:
            page_text.append(page.extract_text() or "")
        except Exception as exc:
            raise ParserError("failed to extract pdf text") from exc

    return "\n\n".join(page_text)


def _extract_docx_text(path):
    try:
        from docx import Document
    except ImportError as exc:
        raise ParserError("python-docx dependency is not installed") from exc

    try:
        document = Document(str(path))
    except Exception as exc:
        raise ParserError("failed to read docx file") from exc

    parts = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text and cell.text.strip()
            ]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)
